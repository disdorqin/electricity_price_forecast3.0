#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Minimal markdown -> docx + pdf converter for the V3.1 research report.
Handles: # headings, **bold**, paragraphs, and pipe tables.
"""
import re, sys, os

SRC = sys.argv[1]
OUTDIR = os.path.dirname(SRC)
BASENAME = os.path.splitext(os.path.basename(SRC))[0]

with open(SRC, "r", encoding="utf-8") as f:
    lines = f.read().split("\n")

# ---- DOCX ----
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def add_table(rows):
    # rows: list of list[str]
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = t.cell(i, j)
            txt = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(8.5)
            if i == 0:
                run.bold = True
    doc.add_paragraph("")

i = 0
buf = []
while i < len(lines):
    ln = lines[i]
    if ln.startswith("# "):
        doc.add_heading(ln[2:].strip(), level=0)
    elif ln.startswith("## "):
        doc.add_heading(ln[3:].strip(), level=1)
    elif ln.startswith("### "):
        doc.add_heading(ln[4:].strip(), level=2)
    elif ln.startswith("#### "):
        doc.add_heading(ln[5:].strip(), level=3)
    elif ln.strip().startswith("|") and ln.strip().endswith("|"):
        # table block
        tbl = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            parts = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            tbl.append(parts)
            i += 1
        # drop separator row if present
        if tbl and all(set(c) <= set("-: ") for c in tbl[1]) and len(tbl) > 1:
            tbl = [tbl[0]] + tbl[2:]
        add_table(tbl)
        continue
    elif ln.strip() == "":
        pass
    else:
        # paragraph, handle **bold**
        txt = ln
        p = doc.add_paragraph()
        # split by **
        segs = re.split(r"(\*\*[^*]+\*\*)", txt)
        for s in segs:
            if s.startswith("**") and s.endswith("**"):
                r = p.add_run(s[2:-2]); r.bold = True
            else:
                p.add_run(s)
    i += 1

docx_path = os.path.join(OUTDIR, BASENAME + ".docx")
doc.save(docx_path)
print("DOCX ->", docx_path)

# ---- PDF (reportlab) ----
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT
import html

pdf_path = os.path.join(OUTDIR, BASENAME + ".pdf")
styles = getSampleStyleSheet()
H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=15, spaceAfter=6)
H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12.5, spaceAfter=4)
H3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, spaceAfter=3)
BODY = ParagraphStyle("BODY", parent=styles["BodyText"], fontSize=9.5, leading=13)
TCELL = ParagraphStyle("TCELL", parent=styles["BodyText"], fontSize=7.5, leading=9)

def esc(s):
    return html.escape(s).replace("**", "")

story = []
i = 0
while i < len(lines):
    ln = lines[i]
    if ln.startswith("# "):
        story.append(Paragraph(esc(ln[2:]), H1))
    elif ln.startswith("## "):
        story.append(Paragraph(esc(ln[3:]), H2))
    elif ln.startswith("### "):
        story.append(Paragraph(esc(ln[4:]), H3))
    elif ln.startswith("#### "):
        story.append(Paragraph(esc(ln[5:]), H3))
    elif ln.strip().startswith("|") and ln.strip().endswith("|"):
        tbl = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            parts = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            tbl.append(parts)
            i += 1
        if tbl and all(set(c) <= set("-: ") for c in tbl[1]) and len(tbl) > 1:
            tbl = [tbl[0]] + tbl[2:]
        ncol = max(len(r) for r in tbl)
        data = [[Paragraph(esc(c), TCELL) for c in (r + [""]*(ncol-len(r)))] for r in tbl]
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.4, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#DCE6F1")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("FONTSIZE", (0,0), (-1,-1), 7.5),
        ]))
        story.append(t)
        story.append(Spacer(1, 4))
        continue
    elif ln.strip() != "":
        story.append(Paragraph(esc(ln), BODY))
        story.append(Spacer(1, 2))
    i += 1

pdf = SimpleDocTemplate(pdf_path, pagesize=A4,
                        leftMargin=18*mm, rightMargin=18*mm,
                        topMargin=16*mm, bottomMargin=16*mm)
pdf.build(story)
print("PDF  ->", pdf_path)
