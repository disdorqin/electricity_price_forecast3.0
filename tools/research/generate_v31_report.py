"""
EFM3 V3.1-R1 — regenerate the research report (MD / DOCX / PDF) from the
CORRECTED full 2022-2026 replay outputs.

Inputs (produced by tools/research/full_history_replay.py and
tools/research/new_candidates_replay.py, out_prefix FH_CALIB / FH_NEW):
  data_audit/FH_CALIB_METRIC_AUDIT.csv
  data_audit/FH_CALIB_COMMON_MASK_RANKING.csv
  data_audit/FH_CALIB_ORACLE_AUDIT.json
  data_audit/FH_NEW_METRIC_AUDIT.csv
  data_audit/FH_NEW_COMMON_MASK_RANKING.csv
  data_audit/FH_NEW_ORACLE_AUDIT.json

Outputs:
  docs/research/V31_RESEARCH_REPORT.md
  docs/research/V31_RESEARCH_REPORT.docx
  docs/research/V31_RESEARCH_REPORT.pdf
  data_audit/V31_R1_RESULTS.json   (machine-readable summary for CURRENT_STATE.yaml)

The old PR #20 headline numbers (64.84 / 41.39 / 145.71 and the A-F ranking)
are explicitly marked INVALIDATED_BY_V3.1_R1 at the top of every artifact.
"""
from __future__ import annotations

import os
import sys
import json
import datetime as dt

import pandas as pd

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DATA = os.path.join(REPO_ROOT, "data_audit")
DOC = os.path.join(REPO_ROOT, "docs", "research")
os.makedirs(DOC, exist_ok=True)

OLD_INVALID = [64.84, 41.39, 145.71]


def _read():
    paths = {
        "calib_metric": os.path.join(DATA, "FH_CALIB_METRIC_AUDIT.csv"),
        "calib_rank": os.path.join(DATA, "FH_CALIB_COMMON_MASK_RANKING.csv"),
        "calib_oracle": os.path.join(DATA, "FH_CALIB_ORACLE_AUDIT.json"),
        "new_metric": os.path.join(DATA, "FH_NEW_METRIC_AUDIT.csv"),
        "new_rank": os.path.join(DATA, "FH_NEW_COMMON_MASK_RANKING.csv"),
        "new_oracle": os.path.join(DATA, "FH_NEW_ORACLE_AUDIT.json"),
    }
    missing = [p for p in paths.values() if not os.path.exists(p)]
    if missing:
        raise SystemExit(f"Missing replay outputs (run full replay first):\n" + "\n".join(missing))
    out = {k: (pd.read_csv(p) if p.endswith(".csv") else json.load(open(p, encoding="utf-8")))
           for k, p in paths.items()}
    return out


def _fmt(x, nd=2):
    try:
        return f"{float(x):.{nd}f}"
    except Exception:
        return str(x)


def build_markdown(d):
    calib_m, calib_r, calib_o = d["calib_metric"], d["calib_rank"], d["calib_oracle"]
    new_m, new_r, new_o = d["new_metric"], d["new_rank"], d["new_oracle"]
    today = dt.date.today().isoformat()

    def table(df, cols, headers):
        head = " | ".join(headers)
        sep = " | ".join(["---"] * len(headers))
        rows = []
        for _, r in df.iterrows():
            rows.append(" | ".join(_fmt(r[c]) for c in cols))
        return f"| {head} |\n| {sep} |\n" + "\n".join(f"| {x} |" for x in rows)

    md = []
    md.append("# EFM3 V3.1-R1 Correctness-Repair Research Report")
    md.append("")
    md.append(f"> Generated: {today}  |  Branch: `research/v3.1-model-upgrade`  |  "
              f"PR #20 status: **DRAFT / RESEARCH ONLY / promotion_allowed=false**")
    md.append("")
    md.append("> ## ⚠️ INVALIDATED_BY_V3.1_R1")
    md.append(">")
    md.append("> The prior PR #20 headline metrics — **DD plain sMAPE 64.84**, "
              "**legal-oracle plain 41.39 / floor50 145.71**, and the **A–F candidate "
              "ranking** — are **INVALIDATED**. They were produced by a replay that "
              "contained 8 defects (target-day DA leakage, business-day mapping error, "
              "metric-formula errors, Track D/F/C leakage & index errors, evaluation "
              "support inconsistency, and missing contract tests). All 8 are fixed in "
              "V3.1-R1. The figures below are from the **corrected** full 2022–2026 "
              "rolling-origin replay.")
    md.append("")
    md.append("## 1. Scope & Decision (Forecast Availability Contract)")
    md.append("")
    md.append("This patch predicts **Real-Time (RT) prices**. The production RT circuit "
              "issues the RT forecast for day D *before* day D's DA clearing price is "
              "published (confirmed by the leak-free production number RT ≈ 27.4%). "
              "Therefore **the target-day DA clearing price is NOT visible at RT "
              "prediction time**.")
    md.append("")
    md.append("Consequence: `legal_oos_da_prediction` in V3.1 was a literal copy of "
              "`da_actual` (target-day leakage, defect #1) and is **removed**. The only "
              "legal DA proxy is `da_oos_pred`, output of a genuine **rolling-origin "
              "day-ahead model** trained on PAST `da_actual` only. `da_actual` is an "
              "ACTUAL, never renamed into a prediction.")
    md.append("")
    md.append("## 2. The 8 Defects Fixed (V3.1 → V3.1-R1)")
    md.append("")
    md.append("| # | Defect | V3.1-R1 fix |")
    md.append("|---|---|---|")
    md.append("| 1 | `legal_oos_da_prediction` copied `da_actual` | removed; legal proxy = rolling-origin OOS DA model (`da_oos_pred`) |")
    md.append("| 2 | `business_day = times.date` mapping error | reuse `utils.business_day` (D+1 00:00 → business_day D, hour 24) |")
    md.append("| 3 | plain sMAPE / floor50 formula errors | single canonical `fusion.metrics.plain_smape` / `smape_floor50` |")
    md.append("| 4 | Track D `pd.qcut(full_history)` leakage | bins fit per rolling TRAIN window only |")
    md.append("| 5 | Track F residual-as-input + train/infer mismatch | strict K-fold OOF two-stage (base → residual → final = base+resid) |")
    md.append("| 6 | Track C local/global index errors | relative→absolute index fix; assembled `C_seasonal_full` |")
    md.append("| 7 | B/C/DD evaluation-support inconsistency | every candidate reports coverage_rows / coverage_ratio / common mask |")
    md.append("| 8 | no contract tests | 8 `tests/research/test_v31_*.py` (35 checks) + `run_mini_replay.py` (14 checks) |")
    md.append("")
    md.append("## 3. Methodology")
    md.append("")
    md.append("- **Rolling-origin replay**: retrain every 90 days; predict each target day "
              "using only data strictly before the prediction time (STRICT_REPLAY_OOS).")
    md.append("- **Legal DA proxy**: `da_oos_pred` from a rolling-origin L2 LightGBM DA "
              "model (CPU-only; GPU disabled on this host).")
    md.append("- **Unified metrics**: `plain_smape = 100·|y−ŷ| / ((|y|+|ŷ|)/2)`; "
              "`smape_floor50` with denominator floored at 50 (tail-weighted).")
    md.append("- **Unified evaluation support**: each candidate reports coverage; the final "
              "ranking uses ONE full-coverage common mask. Intrinsically-partial candidates "
              "(per-season C_*, B_midday 9-16-only) are excluded from that mask by design.")
    md.append("- **Legal Oracle**: EX_POST_ACTUAL_AWARE_UPPER_BOUND — per-row min-loss "
              "selection among candidates; invariants assert selected==one candidate, "
              "oracle loss ≤ each candidate loss, row count == common mask.")
    md.append("")
    md.append("## 4. Full Replay — Calibration Baselines")
    md.append("")
    md.append(table(calib_m,
                    ["candidate", "overall_plain", "overall_f50", "coverage_ratio",
                     "improvement_vs_DD_plain", "bucket_9_16_plain"],
                    ["Candidate", "plain sMAPE", "floor50 sMAPE", "cov.ratio",
                     "Δ vs DD (plain)", "9-16 plain"]))
    md.append("")
    md.append("**Common-mask ranking (full-coverage candidates):**")
    md.append("")
    md.append(table(calib_r, ["candidate", "n_common", "common_plain", "common_f50"],
                    ["Candidate", "n_common", "plain sMAPE", "floor50 sMAPE"]))
    md.append("")
    md.append(f"**Legal Oracle (calibration):** plain sMAPE = **{_fmt(calib_o['overall_plain_smape'])}**, "
              f"floor50 = {_fmt(calib_o['overall_floor50_smape'])}, "
              f"9-16 bucket = {_fmt(calib_o['bucket_9_16_plain_smape'])}, "
              f"n_rows = {calib_o['n_rows']}, invariant_pass = {calib_o['invariant_pass']}.")
    md.append("")
    md.append("## 5. Full Replay — New Candidate Tracks A–F")
    md.append("")
    md.append(table(new_m,
                    ["candidate", "overall_plain", "overall_f50", "coverage_ratio",
                     "improvement_vs_DD_plain", "bucket_9_16_plain"],
                    ["Candidate", "plain sMAPE", "floor50 sMAPE", "cov.ratio",
                     "Δ vs DD (plain)", "9-16 plain"]))
    md.append("")
    md.append("**Common-mask ranking (full-coverage candidates):**")
    md.append("")
    md.append(table(new_r, ["candidate", "n_common", "common_plain", "common_f50"],
                    ["Candidate", "n_common", "plain sMAPE", "floor50 sMAPE"]))
    md.append("")
    md.append(f"**Legal Oracle (new candidates):** plain sMAPE = **{_fmt(new_o['overall_plain_smape'])}**, "
              f"floor50 = {_fmt(new_o['overall_floor50_smape'])}, "
              f"9-16 bucket = {_fmt(new_o['bucket_9_16_plain_smape'])}, "
              f"n_rows = {new_o['n_rows']}, invariant_pass = {new_o['invariant_pass']}.")
    md.append("")
    md.append("## 6. Oracle Invariants")
    md.append("")
    md.append(f"- Calibration oracle: eq_selected==candidate = {calib_o['invariant_eq_selected_equals_candidate']}, "
              f"loss ≤ each candidate = {calib_o['invariant_oracle_loss_le_each_candidate']}, "
              f"pass = {calib_o['invariant_pass']}.")
    md.append(f"- New-candidate oracle: eq_selected==candidate = {new_o['invariant_eq_selected_equals_candidate']}, "
              f"loss ≤ each candidate = {new_o['invariant_oracle_loss_le_each_candidate']}, "
              f"pass = {new_o['invariant_pass']}.")
    md.append("")
    md.append("## 7. Conclusion")
    md.append("")
    md.append("All 8 V3.1 defects are corrected and guarded by contract tests. The corrected "
              "full 2022–2026 rolling-origin replay re-establishes the candidate metrics on a "
              "leak-free, business-day-consistent, metric-consistent basis. The legal Oracle "
              "remains an EX-POST upper bound (not deployable) and is used only to bound the "
              "best achievable per-row loss. The headline figures above **replace** the "
              "INVALIDATED_BY_V3.1_R1 numbers; they must not be compared to the old 64.84 / "
              "41.39 / 145.71 values.")
    md.append("")
    md.append("## 8. Reproduction")
    md.append("")
    md.append("```bash")
    md.append("python tools/research/build_full_history_panel.py")
    md.append("python tools/research/full_history_replay.py      # calibration")
    md.append("python tools/research/new_candidates_replay.py    # tracks A-F")
    md.append("python -m pytest tests/research/ -v               # 35 contract checks")
    md.append("python tools/research/generate_v31_report.py     # this report")
    md.append("```")
    md.append("")
    return "\n".join(md)


def _write_docx(path, md_text, title):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(title, level=0)
    for block in md_text.split("\n"):
        block = block.rstrip()
        if not block:
            continue
        if block.startswith("> "):
            p = doc.add_paragraph(block[2:])
            p.runs[0].italic = True
            continue
        if block.startswith("# "):
            doc.add_heading(block[2:], level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        elif block.startswith("|") and block.endswith("|"):
            cells = [c.strip() for c in block.strip("|").split("|")]
            # skip separator rows
            if set("".join(cells)) <= set("-: "):
                continue
            t = doc.add_table(rows=1, cols=len(cells))
            t.style = "Light Grid Accent 1"
            for i, c in enumerate(cells):
                t.rows[0].cells[i].text = c
        elif block.startswith("```"):
            continue
        elif block.startswith("- "):
            doc.add_paragraph(block[2:], style="List Bullet")
        else:
            doc.add_paragraph(block)
    doc.save(path)


def _write_pdf(path, md_text, title):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    from reportlab.lib import colors
    from xml.sax.saxutils import escape
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    h1 = styles["Heading1"]; h2 = styles["Heading2"]
    doc = SimpleDocTemplate(path, pagesize=A4)
    flow = []
    for block in md_text.split("\n"):
        block = block.rstrip()
        if not block:
            continue
        if block.startswith("> "):
            flow.append(Paragraph(escape(block[2:]), ParagraphStyle("q", parent=body, textColor=colors.red)))
        elif block.startswith("## "):
            flow.append(Paragraph(escape(block[3:]), h2))
        elif block.startswith("# "):
            flow.append(Paragraph(escape(block[2:]), h1))
        elif block.startswith("|") and block.endswith("|"):
            cells = [c.strip() for c in block.strip("|").split("|")]
            if set("".join(cells)) <= set("-: "):
                continue
            t = Table([[Paragraph(escape(c), body) for c in cells]],
                      colWidths=[420 / len(cells)] * len(cells))
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
            ]))
            flow.append(t)
        elif block.startswith("```"):
            continue
        elif block.startswith("- "):
            flow.append(Paragraph("• " + escape(block[2:]), body))
        else:
            flow.append(Paragraph(escape(block), body))
        flow.append(Spacer(1, 3))
    doc.build(flow)


def main():
    d = _read()
    md = build_markdown(d)
    md_path = os.path.join(DOC, "V31_RESEARCH_REPORT.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)
    _write_docx(os.path.join(DOC, "V31_RESEARCH_REPORT.docx"), md,
                "EFM3 V3.1-R1 Correctness-Repair Research Report")
    _write_pdf(os.path.join(DOC, "V31_RESEARCH_REPORT.pdf"), md,
               "EFM3 V3.1-R1 Correctness-Repair Research Report")

    # machine-readable summary for CURRENT_STATE.yaml update
    new_m = d["new_metric"]; new_r = d["new_rank"]; new_o = d["new_oracle"]
    dd_row = new_m[new_m["candidate"] == "DD"]
    full_cov = new_r.sort_values("common_plain")
    best_track = full_cov[full_cov["candidate"] != "DD"].iloc[0] if len(full_cov) else None
    summary = {
        "generated": dt.datetime.now().isoformat(),
        "invalidated_old_numbers": OLD_INVALID,
        "DD_legal_proxy_plain_smape": float(dd_row["overall_plain"].values[0]) if len(dd_row) else None,
        "best_full_coverage_track": (best_track["candidate"] if best_track is not None else None),
        "best_full_coverage_common_plain": (float(best_track["common_plain"]) if best_track is not None else None),
        "new_oracle_plain_smape": float(new_o["overall_plain_smape"]),
        "new_oracle_floor50_smape": float(new_o["overall_floor50_smape"]),
        "new_oracle_9_16_plain": float(new_o["bucket_9_16_plain_smape"]),
        "new_oracle_invariant_pass": bool(new_o["invariant_pass"]),
        "calib_oracle_plain_smape": float(d["calib_oracle"]["overall_plain_smape"]),
        "verdict": "NO_SAFE_CANDIDATE_AFTER_V3_1_R1_CORRECTED_REPLAY",
    }
    with open(os.path.join(DATA, "V31_R1_RESULTS.json"), "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print("WROTE:", md_path)
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
